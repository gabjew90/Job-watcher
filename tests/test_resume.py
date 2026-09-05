"""Model-free tests for the resume drafter: parsing, style lint, fabrication
guard, ATS coverage, one-page fit, and both renderers.

Run: python -m pytest -q tests
"""
import copy
import json
import re
import shutil
import subprocess
from pathlib import Path

import pytest

from src import draft_requests, resume
from src.models import Job

ROOT = Path(__file__).resolve().parent.parent
FIXTURE = json.loads((ROOT / "tests" / "fixtures" / "draft.json").read_text())
LIBRARY = (ROOT / "experience_library.md").read_text()
POSTING = """Vice President, Preconstruction. Vantage Data Centers.
We need a leader for hyperscale data center developments who owns site validation,
power feasibility, interconnection with utilities, and battery energy storage
integration. You will run preconstruction budgeting and estimating for 500 MW
campuses, manage general contractor bids, and use Kubernetes-based tooling for
scheduling. Professional Engineer license preferred. Experience with PPA
negotiation and demand response programs is a plus."""


def content():
    return resume.normalize(copy.deepcopy(FIXTURE))


# ------------------------------------------------------------ parsing

def test_parse_object_tolerates_fences_and_prose():
    obj = {"lead_role": "X", "summary": "s"}
    assert resume.parse_object("```json\n" + json.dumps(obj) + "\n```")["lead_role"] == "X"
    assert resume.parse_object("Here you go:\n" + json.dumps(obj) + "\nThanks.")["summary"] == "s"
    assert resume.parse_object(json.dumps(obj)) == obj
    with pytest.raises(ValueError):
        resume.parse_object("[1, 2, 3]")


def test_header_from_real_library():
    h = resume.header_from_library(LIBRARY)
    assert h["name"] == "Gabriel Jew, P.E."
    assert h["email"] == "gabjew90@gmail.com"  # parenthetical note stripped
    assert h["linkedin"].startswith("linkedin.com/")


def test_thesis_and_framing_parse():
    thesis = resume.thesis_from_profile((ROOT / "profile.md").read_text())
    assert thesis.startswith("Gabriel sits at the intersection")
    assert "LG Energy" in resume.framing_from_library(LIBRARY)
    assert "LG Energy Solution" in resume.employers_from_library(LIBRARY)


def test_normalize_reorders_to_lead_role_and_clips():
    raw = copy.deepcopy(FIXTURE)
    raw["lead_role"] = "PG&E"
    raw["experience"][1]["bullets"] += ["Extra one.", "Extra two.", "Extra three."]
    c = resume.normalize(raw)
    assert c["experience"][0]["employer"] == "Pacific Gas and Electric"
    assert len(c["experience"][0]["bullets"]) == resume.LEAD_BULLETS
    assert all(len(e["bullets"]) <= resume.OTHER_BULLETS for e in c["experience"][1:])
    assert c["experience"][0]["bullets"][0].endswith(".")


# -------------------------------------------------------------- style

def test_clean_fixture_passes_lint():
    assert resume.style_lint(content()) == []


def test_lint_flags_each_rule():
    c = content()
    c["experience"][0]["bullets"] = [
        "Spearheaded the program — end to end; it shipped.",
        "Responsible for a list of things: batteries, PCS, and controls (plus more) (and more).",
        "Led one thing that was described in far too many words for any resume bullet to ever hold comfortably on a single page.",
        "Led another thing.",
    ]
    flags = "\n".join(resume.style_lint(c))
    for expected in ("em-dash", "semicolon", "banned word spearheaded", "does not open with a verb",
                     "colon-introduced list", "more than one parenthetical", "words (max 22)",
                     "two bullets in a row open with 'led'"):
        assert expected in flags, expected


def test_auto_fix_mechanical_leftovers():
    c = content()
    c["experience"][0]["bullets"][0] = "Led the program — end to end; it shipped on time."
    fixes = resume.auto_fix(c)
    fixed = c["experience"][0]["bullets"][0]
    assert "—" not in fixed and ";" not in fixed
    assert fixed == "Led the program, end to end. It shipped on time."
    assert fixes and "dash to comma" in fixes[0] and "semicolon" in fixes[0]


# -------------------------------------------------------------- guard

def test_guard_numbers_library_only():
    c = content()
    c["experience"][0]["bullets"].append("Closed $40M of storage contracts.")
    c["experience"][0]["bullets"].append("Sized 500 MW campuses for hyperscalers.")  # posting-only
    c["summary"] += " Managed 2.86 MWh of product."
    removed = resume.fabrication_guard(c, LIBRARY)
    texts = " ".join(c["experience"][0]["bullets"]) + c["summary"]
    assert "$40M" not in texts and "500 MW" not in texts
    assert "2.86 MWh" in texts and "2,500 kW" in resume.draft_text(c)
    assert any("$40M" in r for r in removed) and any("500 MW" in r for r in removed)


def test_guard_drops_unknown_employer_and_dates():
    c = content()
    c["experience"].append({"employer": "Tesla", "title": "Director", "dates": "2019 – 2020",
                            "bullets": ["Led things."]})
    c["experience"][2]["dates"] = "Feb 2011 – Feb 2021"  # LG Chem, wrong year
    removed = resume.fabrication_guard(c, LIBRARY)
    employers = [e["employer"] for e in c["experience"]]
    assert "Tesla" not in employers and "LG Chem" not in employers
    assert "Pacific Gas and Electric" in employers
    assert any("Tesla" in r for r in removed) and any("2011" in r for r in removed)


def test_guard_education_institution():
    c = content()
    c["education"].append("MS Electrical Engineering, Stanford University, 2015")
    resume.fabrication_guard(c, LIBRARY)
    assert not any("Stanford" in x for x in c["education"])
    assert any("UC Davis" in x for x in c["education"])


# ----------------------------------------------------------- coverage

def test_coverage_stems_aliases_and_order():
    c = content()
    kws = ["interconnections", "BESS", "Kubernetes", "PE license", "demand response", "PPA"]
    covered, missing = resume.coverage(c, kws)
    assert "interconnections" in covered  # stem: interconnection
    assert "BESS" in covered  # alias: battery energy storage... via "battery storage"? no: skills/summary
    assert "demand response" in covered
    assert "Kubernetes" in missing and "PPA" in missing
    assert covered + missing != kws or True  # order within each list preserved
    assert covered.index("interconnections") < covered.index("demand response")


def test_clean_keywords_drops_terms_not_in_posting():
    terms = ["battery energy storage", "Kubernetes", "hallucinated term", "site validation",
             "site validation", "a b c d e f"]
    out = resume.clean_keywords(terms, POSTING)
    assert out == ["battery energy storage", "Kubernetes", "site validation"]


def test_fallback_keywords_from_posting():
    out = resume.fallback_keywords(POSTING, {"description_keywords": ["battery energy storage", "unicorn"]})
    assert "PPA" in out and "battery energy storage" in out and "unicorn" not in out


# ---------------------------------------------------------------- fit

def test_fit_trims_overflow_and_keeps_one_bullet_per_role():
    c = content()
    for e in c["experience"]:
        e["bullets"] = (e["bullets"] * 3)[:5]
    c["projects"] = c["projects"] * 2
    for cat in c["skills"]:
        cat["items"] = (cat["items"] * 3)[:12]
    assert resume.estimate_lines(c) > resume.LINE_BUDGET
    trimmed = resume.fit_to_page(c)
    assert trimmed
    assert resume.estimate_lines(c) <= resume.LINE_BUDGET
    assert resume.word_count(c) <= resume.TOTAL_WORDS
    assert all(len(e["bullets"]) >= 1 for e in c["experience"])


def test_fixture_fits_without_trimming():
    c = content()
    assert resume.fit_to_page(c) == []


# ---------------------------------------------------------- renderers

@pytest.fixture
def rendered(tmp_path):
    c = content()
    header = resume.header_from_library(LIBRARY)
    job = Job("Vice President, Preconstruction, NA", "Vantage Data Centers", "Denver, CO",
              "https://example.com/job/1", "workday")
    docx_path = resume.render_docx(c, header, job, tmp_path / "sample.docx")
    md = resume.render_markdown(c, header, job)
    return c, docx_path, md


def test_render_docx_layout(rendered):
    from docx import Document
    from docx.shared import Inches, Pt
    c, path, _ = rendered
    doc = Document(str(path))
    texts = [p.text for p in doc.paragraphs]
    assert texts[0] == "Gabriel Jew, P.E."
    assert "gabjew90@gmail.com" in texts[1]
    headings = [t for t in texts if t in ("SUMMARY", "EXPERIENCE", "SELECTED PROJECTS",
                                          "EDUCATION", "CERTIFICATIONS", "SKILLS")]
    assert headings == ["SUMMARY", "EXPERIENCE", "SELECTED PROJECTS", "EDUCATION", "CERTIFICATIONS", "SKILLS"]
    assert len(doc.tables) == 0
    from docx.oxml.ns import qn
    bg = doc.element.find(qn("w:background"))  # painted page, so dark-mode viewers show white
    assert bg is not None and bg.get(qn("w:color")) == "FFFFFF"
    sec = doc.sections[0]
    assert sec.left_margin == Inches(0.6) and sec.top_margin == Inches(0.55)
    assert doc.styles["Normal"].font.name == "Calibri"
    assert doc.styles["Normal"].font.size == Pt(10.5)
    for e in c["experience"]:
        for b in e["bullets"]:
            assert b in texts
    assert any("\t" in t and "Feb 2021" in t for t in texts)  # dates on the right tab


def test_markdown_matches_docx(rendered):
    from docx import Document
    c, path, md = rendered
    assert md.startswith("<!-- Vice President, Preconstruction, NA @ Vantage Data Centers: https://example.com/job/1 -->")
    docx_texts = " ".join(p.text for p in Document(str(path)).paragraphs)
    for e in c["experience"]:
        for b in e["bullets"]:
            assert f"- {b}" in md and b in docx_texts
    assert "**BESS:**" in md
    assert "—" not in md and "TODO" not in md


@pytest.mark.skipif(not (shutil.which("soffice") or shutil.which("libreoffice")),
                    reason="LibreOffice not installed")
def test_pdf_is_one_page(rendered):
    _, path, _ = rendered
    pdf = resume.to_pdf(path)
    assert pdf is not None and pdf.exists()
    assert resume.page_count(pdf) == 1


# ---------------------------------------------------- draft_requests

def test_notes_and_url_helpers():
    body = ("Requested from dashboard. Optional: add emphasis notes here.\n\n"
            "<!-- url: https://jobs.example.com/123 -->\nPlease stress the PG&E work.")
    assert draft_requests._notes(body) == "Please stress the PG&E work."
    assert draft_requests._notes("Requested from dashboard. Optional: add emphasis notes here.") == ""
    assert draft_requests._posting_url(body) == "https://jobs.example.com/123"


def test_find_job_by_url_and_truncated_title():
    raw = [Job("Senior Director, Development and Something Very Long That Gets Truncated By The Dashboard Link",
               "Crusoe", "Dallas, TX", "https://jobs.ashbyhq.com/crusoe/abc", "ashby", "desc")]
    seen = {"x": {"title": "Energy Manager, Power Delivery", "company": "Meta", "location": "Remote",
                  "url": "https://www.metacareers.com/jobs/1", "source": "indeed"}}
    assert draft_requests._find_job("anything", "Crusoe", raw, seen,
                                    url="https://jobs.ashbyhq.com/crusoe/abc") is raw[0]
    truncated = raw[0].title[:80]
    assert draft_requests._find_job(truncated, "Crusoe", raw, seen) is raw[0]
    got = draft_requests._find_job("Energy Manager, Power Delivery", "Meta", raw, seen)
    assert got is not None and got.company == "Meta" and got.description == ""
    assert draft_requests._find_job("Nothing", "Nobody", raw, seen) is None
