"""On-demand resume drafting: plain prose, code-owned one-page layout.

The model writes words only, as one JSON object (schema in DRAFT_PROMPT).
Everything else is code:

- layout: a fixed one-page DOCX (python-docx), a PDF of that same DOCX via
  LibreOffice, and a Markdown twin with identical wording for GitHub preview
- style: resume_style.md (owner-curated) is shown to the model, and
  style_lint() checks the result; violations go back to the revision pass,
  mechanical leftovers are fixed in code, the rest are reported
- truth: fabrication_guard() drops any unit whose numbers, years or
  employers do not appear in experience_library.md (drop, never rewrite)
- ATS: keyword extraction from the posting, coverage check against the
  draft, one library-bounded revision pass to close what the library can
- length: fit_to_page() trims by priority until the estimate fits, then the
  PDF page count is the ground truth

Gaps (requirements the library cannot evidence) never enter the resume;
they are returned for the issue comment.
"""
import json
import logging
import math
import os
import re
import shutil
import subprocess
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path

from . import triage
from .models import Job
from .util import company_key, role_excerpt

log = logging.getLogger(__name__)

LIBRARY = Path("experience_library.md")
PROFILE = Path("profile.md")
STYLE = Path("resume_style.md")
DRAFTS_DIR = Path("drafts")

DRAFT_MODEL = os.environ.get("JOBWATCH_DRAFT_MODEL", "claude-sonnet-5")
KEYWORD_MODEL = os.environ.get("JOBWATCH_KEYWORD_MODEL", "claude-haiku-4-5-20251001")

# Budgets. The layout is Letter, 0.6 in side margins, Calibri 10.5 pt:
# about 58 lines of body text; 54 leaves room for spacing drift.
SUMMARY_WORDS = 60
SENTENCE_WORDS = 30
LEAD_BULLETS = 5
OTHER_BULLETS = 3
BULLET_WORDS = 22
MAX_PROJECTS = 3
MAX_SKILL_CATEGORIES = 4
MAX_SKILL_ITEMS = 12
# Calibrated on a rendered page: 407 words / ~62 estimated lines filled
# about 85% of the sheet. The PDF page count is the final judge anyway.
TOTAL_WORDS = 560
LINE_BUDGET = 62
CHARS_PER_LINE = 115
MAX_PDF_ROUNDS = 3

SECTION_ORDER = ("summary", "experience", "projects", "education", "certifications", "skills")


# ---------------------------------------------------------------- prompts

SCHEMA = """{
  "lead_role": "<employer name exactly as in the library: the role to list first>",
  "summary": "<2-3 sentences, 60 words max>",
  "experience": [
    {"employer": "<as in library>", "title": "<as in library>", "dates": "<as in library>",
     "bullets": ["<verb first, 22 words max>"]}
  ],
  "projects": [{"name": "<as in library>", "line": "<one sentence, 22 words max>"}],
  "education": ["<one line each>"],
  "certifications": ["<one line each>"],
  "skills": [{"category": "<name>", "items": ["<term>"]}],
  "gaps": ["<must-have posting requirement the library cannot evidence>"],
  "omitted_requirements": ["<nice-to-have posting ask the library cannot evidence>"]
}"""

HARD_RULES = """HARD RULES
- Draw ONLY from the experience library. Never invent metrics, numbers,
  skills, employers, dates, or accomplishments that are not written there.
  A number that is not in the library must not appear anywhere.
- Never write TODO or placeholder text. If the posting asks for something
  the library cannot evidence, list it under "gaps" (must-haves) or
  "omitted_requirements" (nice-to-haves) and leave it out of the resume.
- Do not write the header (name, contact details): the code adds it.
- Prefer the library's own wording. Change words only to match the
  posting's terms or to follow the style guide.
- Lead with the role the framing guidance names for this kind of posting
  (set "lead_role" to that employer); up to 5 bullets for it, up to 3 for
  every other role. Up to 3 projects, only if relevant. Up to 4 skill
  categories with up to 12 items each, posting terms first.
- No first person. Return ONLY the JSON object, no prose, no fences."""

DRAFT_PROMPT = """Write the content of a tailored one-page resume for the job posting below,
as one JSON object in exactly this shape:

{schema}

{hard_rules}

STYLE GUIDE (follow it exactly):
{style}
{notes}
CANDIDATE THESIS (how the candidate should come across):
{thesis}

FRAMING GUIDANCE (which current role to lead with):
{framing}

JOB POSTING:
{job}

EXPERIENCE LIBRARY (the only source of facts):
{library}"""

KEYWORD_PROMPT = """List the terms an applicant-tracking system would match on for this job
posting: hard skills, tools and software, domain nouns, certifications and
licenses, methods. 15 to 25 terms, each 4 words or fewer, spelled exactly as
the posting spells them. Exclude soft skills, company names, benefits, and
legal boilerplate. Return ONLY a JSON array of strings.

JOB POSTING:
{job}"""

REVISE_PROMPT = """Revise the resume content below. Keep the same JSON shape and return ONLY
the JSON object, plus one extra key "revision_notes": a list of one-line
strings, one per ATS term you left out, with the reason.

Do these things and nothing else:
1. ATS terms missing from the draft: {missing}
   Work a term in ONLY where the experience library genuinely supports it:
   reword an existing bullet to use the posting's term, or add the term to
   the matching skills category. If the library does not support a term,
   leave it out and say so in revision_notes.
2. Style violations to fix: {violations}
3. Keep every bullet count, word budget and section as it is. Never add or
   change a number, employer, date, or the lead_role.

{hard_rules}

STYLE GUIDE:
{style}

EXPERIENCE LIBRARY (the only source of facts):
{library}

CURRENT DRAFT:
{draft}"""


# ------------------------------------------------------- library parsing

def _section(text: str, heading: str) -> str:
    m = re.search(rf"^## {re.escape(heading)}[^\n]*\n(.*?)(?=^## |\Z)", text, re.S | re.M)
    return m.group(1).strip() if m else ""


def header_from_library(text: str) -> dict:
    """The contact header from '## Header' lines like '- Email: x (note)'."""
    out = {}
    for line in _section(text, "Header").splitlines():
        m = re.match(r"\s*-\s*([A-Za-z]+):\s*(.+)", line)
        if m:
            out[m.group(1).lower()] = m.group(2).split(" (")[0].strip()
    return out


def framing_from_library(text: str) -> str:
    return _section(text, "FRAMING GUIDANCE FOR DRAFTS")


def thesis_from_profile(text: str) -> str:
    m = re.search(r"CORE THESIS:(.*?)(?:\n\s*\n|\Z)", text, re.S)
    return re.sub(r"\s+", " ", m.group(1)).strip() if m else ""


def employers_from_library(text: str) -> list[str]:
    return [m.strip() for m in re.findall(r"^### (.+?) — ", text, re.M)]


# ------------------------------------------------------------- parsing

def parse_object(text: str) -> dict:
    """The model's JSON object, tolerant of fences and surrounding prose."""
    text = re.sub(r"^```(?:json)?\s*|\s*```$", "", text.strip(), flags=re.S)
    start, end = text.find("{"), text.rfind("}")
    if start == -1 or end == -1 or end < start:
        raise ValueError(f"no JSON object in response: {text[:200]}")
    return json.loads(text[start:end + 1])


def _s(value) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _clean_bullet(text: str) -> str:
    text = re.sub(r"^[\-•\*\d\.\)\s]+", "", _s(text))
    if text and text[-1] not in ".!?":
        text += "."
    return text


def normalize(content: dict) -> dict:
    """Coerce the model's object into the schema, clip to budgets, and put
    the lead role first."""
    out = {
        "lead_role": _s(content.get("lead_role")),
        "summary": _s(content.get("summary")),
        "experience": [],
        "projects": [],
        "education": [_s(x) for x in content.get("education") or [] if _s(x)],
        "certifications": [_s(x) for x in content.get("certifications") or [] if _s(x)],
        "skills": [],
        "gaps": [_s(x) for x in content.get("gaps") or [] if _s(x)],
        "omitted_requirements": [_s(x) for x in content.get("omitted_requirements") or [] if _s(x)],
        "revision_notes": [_s(x) for x in content.get("revision_notes") or [] if _s(x)],
    }
    for e in content.get("experience") or []:
        if not isinstance(e, dict):
            continue
        bullets = [_clean_bullet(b) for b in (e.get("bullets") or []) if _s(b)]
        out["experience"].append({"employer": _s(e.get("employer")), "title": _s(e.get("title")),
                                  "dates": _s(e.get("dates")), "bullets": bullets})
    for p in (content.get("projects") or [])[:MAX_PROJECTS]:
        if isinstance(p, dict) and _s(p.get("name")):
            out["projects"].append({"name": _s(p.get("name")), "line": _clean_bullet(p.get("line"))})
    for c in (content.get("skills") or [])[:MAX_SKILL_CATEGORIES]:
        if isinstance(c, dict) and _s(c.get("category")):
            items = [_s(i).rstrip(".") for i in (c.get("items") or []) if _s(i)][:MAX_SKILL_ITEMS]
            if items:
                out["skills"].append({"category": _s(c.get("category")), "items": items})
    # Lead role first, then clip bullet counts.
    lead = company_key(out["lead_role"])
    if lead:
        for i, e in enumerate(out["experience"]):
            if company_key(e["employer"]) == lead:
                out["experience"].insert(0, out["experience"].pop(i))
                break
    for i, e in enumerate(out["experience"]):
        e["bullets"] = e["bullets"][:LEAD_BULLETS if i == 0 else OTHER_BULLETS]
    out["experience"] = [e for e in out["experience"] if e["employer"]]
    return out


# ---------------------------------------------------------------- style

BANNED_WORDS = (
    "spearheaded", "leveraged", "leveraging", "leverage", "drive", "driving", "driven",
    "cross-functional", "cross functional", "passionate", "proven track record", "robust",
    "seamless", "seamlessly", "holistic", "dynamic", "cutting-edge", "synergy", "synergies",
    "utilize", "utilizing", "utilized", "impactful", "innovative", "world-class",
    "best-in-class", "stakeholders", "ecosystem", "transformative", "empower", "empowering",
    "at the intersection of", "seasoned", "results-driven", "results-oriented",
)

VERBS = set("""
lead leads led manage manages managed run runs ran build builds built design designs designed
launch launches launched negotiate negotiates negotiated win wins won cut cuts size sized sizes
model models modeled write writes wrote ship ships shipped own owns owned set sets develop
develops developed deliver delivers delivered author authors authored present presents presented
coordinate coordinates coordinated direct directs directed drive plan plans planned price priced
prices secure secures secured create creates created define defines defined prioritize
prioritizes prioritized analyze analyzes analyzed evaluate evaluates evaluated improve improves
improved reduce reduces reduced increase increases increased grow grows grew scale scales scaled
support supports supported advise advises advised research researches researched test tests tested
validate validates validated integrate integrates integrated implement implements implemented
establish establishes established form forms formed found founded oversee oversees oversaw
maintain maintains maintained operate operates operated engineer engineers engineered select
selects selected specify specifies specified schedule schedules scheduled train trains trained
recruit recruits recruited hire hires hired mentor mentors mentored partner partners partnered
sell sells sold source sources sourced procure procures procured structure structures structured
close closes closed originate originates originated forecast forecasts forecasted optimize
optimizes optimized automate automates automated deploy deploys deployed publish publishes
published report reports reported track tracks tracked audit audits audited assess assesses
assessed review reviews reviewed approve approves approved beat beats exceed exceeds exceeded
generate generates generated produce produces produced save saves saved handle handles handled
resolve resolves resolved translate translates translated bridge bridges bridged quantify
quantifies quantified compile compiles compiled gather gathers gathered interview interviews
interviewed survey surveys surveyed map maps mapped shape shapes shaped steer steers steered
guide guides guided champion champions championed sponsor sponsors sponsored serve serves served
help helps helped provide provides provided perform performs performed conduct conducts conducted
execute executes executed complete completes completed achieve achieves achieved earn earns earned
convert converts converted turn turns turned bring brings brought take takes took make makes made
give gives gave keep keeps kept hold holds held drove built
""".split())


def _sentences(text: str) -> list[str]:
    return [s for s in re.split(r"(?<=[.!?])\s+", text.strip()) if s]


def _words(text: str) -> int:
    return len(re.findall(r"\S+", text))


def _units(content: dict) -> list[tuple[str, str]]:
    """(label, text) for every prose unit the lint and guard look at."""
    units = [(f"summary sentence {i + 1}", s) for i, s in enumerate(_sentences(content.get("summary", "")))]
    for e in content.get("experience", []):
        for j, b in enumerate(e.get("bullets", [])):
            units.append((f"{e['employer']} bullet {j + 1}", b))
    for p in content.get("projects", []):
        units.append((f"project {p['name']}", p["line"]))
    for line in content.get("education", []):
        units.append(("education", line))
    for line in content.get("certifications", []):
        units.append(("certification", line))
    return units


def style_lint(content: dict) -> list[str]:
    """Violations of the code-enforced style rules, as 'label: reason'."""
    flags = []
    banned = sorted(BANNED_WORDS, key=len, reverse=True)
    for label, text in _units(content):
        low = text.lower()
        if "—" in text or " – " in text:
            flags.append(f"{label}: em-dash")
        if ";" in text:
            flags.append(f"{label}: semicolon")
        hits = [w for w in banned if re.search(rf"\b{re.escape(w)}\b", low)]
        if hits:
            flags.append(f"{label}: banned word {', '.join(hits)}")
        if text.count("(") > 1:
            flags.append(f"{label}: more than one parenthetical")
        if label.startswith("summary"):
            if _words(text) > SENTENCE_WORDS:
                flags.append(f"{label}: {_words(text)} words (max {SENTENCE_WORDS})")
        elif "bullet" in label or label.startswith("project"):
            if _words(text) > BULLET_WORDS:
                flags.append(f"{label}: {_words(text)} words (max {BULLET_WORDS})")
            if re.search(r":\s", text):
                flags.append(f"{label}: colon-introduced list")
            first = re.sub(r"[^a-z]", "", low.split(" ", 1)[0]) if "bullet" in label else ""
            if first and first not in VERBS:
                flags.append(f"{label}: does not open with a verb ('{first}')")
    if _words(content.get("summary", "")) > SUMMARY_WORDS:
        flags.append(f"summary: {_words(content['summary'])} words (max {SUMMARY_WORDS})")
    for e in content.get("experience", []):
        openers = [re.sub(r"[^a-z]", "", b.lower().split(" ", 1)[0]) for b in e.get("bullets", [])]
        for a, b in zip(openers, openers[1:]):
            if a and a == b:
                flags.append(f"{e['employer']}: two bullets in a row open with '{a}'")
                break
    return flags


def _fix_text(text: str) -> tuple[str, list[str]]:
    fixes = []
    if "—" in text or " – " in text:
        text = re.sub(r"\s*—\s*|\s+–\s+", ", ", text)
        fixes.append("dash to comma")
    if ";" in text:
        def _cap(m):
            return ". " + m.group(1).upper()
        text = re.sub(r";\s*(\w)", _cap, text)
        fixes.append("semicolon to full stop")
    text = re.sub(r"\s{2,}", " ", text).replace(" ,", ",").replace(",,", ",")
    return text, fixes


def auto_fix(content: dict) -> list[str]:
    """Mechanical style fixes applied in place; returns what changed."""
    changed = []
    summary, fixes = _fix_text(content.get("summary", ""))
    if fixes:
        content["summary"] = summary
        changed.append(f"summary: {', '.join(fixes)}")
    for e in content.get("experience", []):
        new = []
        for j, b in enumerate(e.get("bullets", [])):
            b2, fixes = _fix_text(b)
            new.append(b2)
            if fixes:
                changed.append(f"{e['employer']} bullet {j + 1}: {', '.join(fixes)}")
        e["bullets"] = new
    for p in content.get("projects", []):
        p["line"], fixes = _fix_text(p["line"])
        if fixes:
            changed.append(f"project {p['name']}: {', '.join(fixes)}")
    return changed


# ---------------------------------------------------------------- guard

NUM_RE = re.compile(r"\$?\d[\d,.]*\+?(?:\s?(?:%|[kKMBG]\b|[kMG]Wh?\b|million|billion))?")
YEAR_RE = re.compile(r"\b(?:19|20)\d{2}\b")
SCHOOL_RE = re.compile(r"university|college|school|institute|\bUC\b|polytechnic", re.I)


def _num_norm(token: str) -> str:
    return token.lower().replace(",", "").replace(" ", "").replace("+", "").rstrip(".")


def _bare(token: str) -> str:
    return re.sub(r"[^\d.]", "", token).rstrip(".")


def _allowed_numbers(library: str) -> tuple[set, set]:
    full, bare = set(), set()
    for tok in NUM_RE.findall(library):
        n = _num_norm(tok)
        full.add(n)
        bare.add(_bare(n))
    return full, bare


def _number_violations(text: str, full: set, bare: set) -> list[str]:
    bad = []
    for tok in NUM_RE.findall(text):
        n = _num_norm(tok)
        if not _bare(n):
            continue
        has_unit = bool(re.search(r"[a-z%$]", n))
        if has_unit and n not in full:
            # A bare-number match with a different unit is still a fabrication.
            bad.append(tok.strip())
        elif not has_unit and _bare(n) not in bare:
            bad.append(tok.strip())
    return bad


def fabrication_guard(content: dict, library: str) -> list[str]:
    """Drop every unit whose numbers, years or employers are not in the
    library. Drop, never rewrite. Returns 'text: reason' per removal."""
    full, bare = _allowed_numbers(library)
    lib_low = re.sub(r"\s+", " ", library.lower())
    lib_years = set(YEAR_RE.findall(library))
    removed = []

    def ok_numbers(text: str, label: str) -> bool:
        bad = _number_violations(text, full, bare)
        if bad:
            removed.append(f"{label} '{text[:90]}': number {', '.join(bad)} not in library")
            return False
        return True

    kept_sentences = [s for s in _sentences(content.get("summary", "")) if ok_numbers(s, "summary")]
    content["summary"] = " ".join(kept_sentences)

    experience = []
    for e in content.get("experience", []):
        emp = re.sub(r"\s+", " ", e["employer"].lower())
        if emp not in lib_low:
            removed.append(f"role '{e['employer']}': employer not in library")
            continue
        years = set(YEAR_RE.findall(e["dates"]))
        if years - lib_years:
            removed.append(f"role '{e['employer']}': dates {e['dates']} not in library")
            continue
        e["bullets"] = [b for b in e["bullets"] if ok_numbers(b, f"{e['employer']} bullet")]
        experience.append(e)
    content["experience"] = experience

    projects = []
    for p in content.get("projects", []):
        if re.sub(r"\s+", " ", p["name"].lower()) not in lib_low:
            removed.append(f"project '{p['name']}': not in library")
        elif ok_numbers(p["line"], f"project {p['name']}"):
            projects.append(p)
    content["projects"] = projects

    education = []
    for line in content.get("education", []):
        if not ok_numbers(line, "education"):
            continue
        parts = [x.strip() for x in re.split(r",| — ", line)]
        schools = [x for x in parts if SCHOOL_RE.search(x)]
        if any(re.sub(r"\s+", " ", s.lower()) not in lib_low for s in schools):
            removed.append(f"education '{line[:90]}': institution not in library")
            continue
        education.append(line)
    content["education"] = education
    content["certifications"] = [c for c in content.get("certifications", [])
                                 if ok_numbers(c, "certification")]
    for cat in content.get("skills", []):
        cat["items"] = [i for i in cat["items"] if ok_numbers(i, f"skills {cat['category']}")]
    content["skills"] = [c for c in content.get("skills", []) if c["items"]]
    return removed


# ------------------------------------------------------------- coverage

ALIASES = {
    "bess": ["battery energy storage", "battery energy storage system"],
    "ess": ["energy storage system"],
    "pcs": ["power conversion system"],
    "ica": ["integration capacity analysis"],
    "nwa": ["non wires alternative", "non wires alternatives"],
    "pe": ["professional engineer"],
    "voc": ["voice of customer", "voice of the customer"],
    "dr": ["demand response"],
    "ai": ["artificial intelligence"],
    "ml": ["machine learning"],
    "ppa": ["power purchase agreement"],
    "hvac": ["heating ventilation and air conditioning"],
}
_EXPANSIONS = {}
for _k, _vs in ALIASES.items():
    for _v in _vs:
        _EXPANSIONS.setdefault(_v, []).append(_k)


def _norm(text: str) -> str:
    t = text.lower().replace("&", " and ")
    t = re.sub(r"\bdatacenters?\b", "data center", t)
    t = re.sub(r"[/\-–—_]", " ", t)
    t = re.sub(r"[^a-z0-9+#\s]", " ", t)
    return re.sub(r"\s+", " ", t).strip()


def _stem(tok: str) -> str:
    if len(tok) <= 3:
        return tok
    if tok.endswith("ies"):
        return tok[:-3] + "y"
    if tok.endswith("ing") and len(tok) > 6:
        return tok[:-3]
    if tok.endswith("ed") and len(tok) > 5:
        return tok[:-2]
    if tok.endswith("es") and len(tok) > 5:
        return tok[:-2]
    if tok.endswith("s") and len(tok) > 4:
        return tok[:-1]
    return tok


def _stems(text: str) -> list[str]:
    return [_stem(t) for t in _norm(text).split()]


def _contains(hay: list[str], needle: list[str]) -> bool:
    n = len(needle)
    return n > 0 and any(hay[i:i + n] == needle for i in range(len(hay) - n + 1))


def draft_text(content: dict) -> str:
    parts = [content.get("summary", "")]
    for e in content.get("experience", []):
        parts += [e["title"], *e["bullets"]]
    parts += [p["line"] for p in content.get("projects", [])]
    parts += content.get("education", []) + content.get("certifications", [])
    for c in content.get("skills", []):
        parts += c["items"]
    return " ".join(parts)


def coverage(content: dict, keywords: list[str]) -> tuple[list[str], list[str]]:
    hay = _stems(draft_text(content))
    covered, missing = [], []
    for term in keywords:
        forms = [term] + ALIASES.get(_norm(term), []) + _EXPANSIONS.get(_norm(term), [])
        (covered if any(_contains(hay, _stems(f)) for f in forms) else missing).append(term)
    return covered, missing


def clean_keywords(terms: list, posting: str) -> list[str]:
    """Dedupe, cap length, and drop anything not actually in the posting."""
    hay = _stems(posting)
    seen, out = set(), []
    for t in terms:
        t = _s(t)
        key = _norm(t)
        if not key or key in seen or len(key.split()) > 4:
            continue
        if not _contains(hay, _stems(t)):
            continue
        seen.add(key)
        out.append(t)
    return out[:25]


def fallback_keywords(posting: str, config: dict | None) -> list[str]:
    terms = sorted(set(re.findall(r"\b[A-Z]{2,6}\b", posting)))
    for key in ("description_keywords", "priority_topics", "title_keywords"):
        terms += (config or {}).get(key, [])
    return clean_keywords(terms, posting)


# ------------------------------------------------------------------ fit

def word_count(content: dict) -> int:
    return _words(draft_text(content)) + sum(_words(e["employer"]) for e in content.get("experience", []))


def estimate_lines(content: dict) -> int:
    lines = 2  # name + contact
    lines += 2 + sum(math.ceil(len(s) / CHARS_PER_LINE) for s in [content.get("summary", "")])
    lines += 2
    for e in content.get("experience", []):
        lines += 1 + sum(math.ceil(len(b) / CHARS_PER_LINE) for b in e["bullets"])
    if content.get("projects"):
        lines += 2 + sum(math.ceil((len(p["name"]) + len(p["line"]) + 3) / CHARS_PER_LINE)
                         for p in content["projects"])
    for key in ("education", "certifications"):
        if content.get(key):
            lines += 2 + len(content[key])
    if content.get("skills"):
        lines += 2 + sum(math.ceil((len(c["category"]) + len(", ".join(c["items"])) + 2) / CHARS_PER_LINE)
                         for c in content["skills"])
    return lines


def _over_budget(content: dict) -> bool:
    return estimate_lines(content) > LINE_BUDGET or word_count(content) > TOTAL_WORDS


def trim_one(content: dict) -> str | None:
    """Remove the lowest-priority item. Returns what was removed, or None."""
    exp = content.get("experience", [])
    projects = content.get("projects", [])
    if len(projects) > 2:
        p = projects.pop()
        return f"project '{p['name']}'"
    # Oldest non-lead role first, cycling toward the newest, never below 1.
    for e in reversed(exp[1:]):
        if len(e["bullets"]) > 1:
            b = e["bullets"].pop()
            return f"{e['employer']} bullet '{b[:60]}'"
    for c in content.get("skills", []):
        if len(c["items"]) > 8:
            c["items"] = c["items"][:8]
            return f"skills {c['category']} beyond 8 items"
    if exp and len(exp[0]["bullets"]) > 3:
        b = exp[0]["bullets"].pop()
        return f"{exp[0]['employer']} bullet '{b[:60]}'"
    if projects:
        p = projects.pop()
        return f"project '{p['name']}'"
    if len(exp) > 2:
        e = exp.pop()
        return f"role '{e['employer']}, {e['title']}'"
    return None


def fit_to_page(content: dict) -> list[str]:
    trimmed = []
    while _over_budget(content):
        removed = trim_one(content)
        if removed is None:
            break
        trimmed.append(removed)
    return trimmed


# ------------------------------------------------------------ renderers

def render_markdown(content: dict, header: dict, job: Job | None = None) -> str:
    lines = []
    if job is not None:
        lines.append(f"<!-- {job.title} @ {job.company}: {job.url} -->\n")
    lines.append(f"# {header.get('name', '')}")
    lines.append(_contact_line(header))
    lines.append("\n## Summary\n" + content.get("summary", ""))
    lines.append("\n## Experience")
    for e in content.get("experience", []):
        lines.append(f"\n### {e['employer']} | {e['title']}\n*{e['dates']}*")
        lines += [f"- {b}" for b in e["bullets"]]
    if content.get("projects"):
        lines.append("\n## Selected projects")
        lines += [f"- **{p['name']}**: {p['line']}" for p in content["projects"]]
    if content.get("education"):
        lines.append("\n## Education")
        lines += [f"- {x}" for x in content["education"]]
    if content.get("certifications"):
        lines.append("\n## Certifications")
        lines += [f"- {x}" for x in content["certifications"]]
    if content.get("skills"):
        lines.append("\n## Skills")
        lines += [f"**{c['category']}:** {', '.join(c['items'])}  " for c in content["skills"]]
    return "\n".join(lines) + "\n"


def _contact_line(header: dict) -> str:
    return " | ".join(x for x in (header.get("location"), header.get("phone"), header.get("email"),
                                  header.get("linkedin"), header.get("github")) if x)


THEMES = {
    # Ruled small-caps headings, employer first. Reads like a well-set
    # traditional resume.
    "classic": dict(font="Calibri", body=10.5, name=16, name_color=None, contact=9.5,
                    contact_color=None, name_rule=False, heading_size=9.5,
                    heading_color="333333", heading_spacing=0, heading_before=7,
                    rule=True, rule_color="999999", rule_size=6, role_order="employer",
                    dates_italic=True, dates_color=None, label_color=None),
    # No rules: letter-spaced slate headings, Arial, title before employer,
    # grey dates. The quiet modern look.
    "modern": dict(font="Arial", body=10, name=18, name_color="1F2933", contact=9,
                   contact_color="5B6470", name_rule=True, heading_size=8.5,
                   heading_color="2F4F6F", heading_spacing=30, heading_before=9,
                   rule=False, rule_color="D9DDE2", rule_size=4, role_order="title",
                   dates_italic=False, dates_color="5B6470", label_color="2F4F6F"),
    # One accent colour on the name, headings and their rules; otherwise
    # classic bones.
    "accent": dict(font="Calibri", body=10.5, name=16, name_color="1B4965", contact=9.5,
                   contact_color="5B6470", name_rule=False, heading_size=9,
                   heading_color="1B4965", heading_spacing=20, heading_before=6,
                   rule=True, rule_color="1B4965", rule_size=8, role_order="employer",
                   dates_italic=True, dates_color="5B6470", label_color="1B4965"),
}
DEFAULT_THEME = os.environ.get("JOBWATCH_RESUME_THEME", "accent")


def render_docx(content: dict, header: dict, job: Job | None, path: Path,
                theme: str = DEFAULT_THEME) -> Path:
    from docx import Document
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    t = THEMES[theme]
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.left_margin = sec.right_margin = Inches(0.6)
    sec.top_margin = sec.bottom_margin = Inches(0.55)

    def _rgb(hex6):
        return RGBColor.from_string(hex6) if hex6 else None

    def _font(style, size, bold=None):
        style.font.name = t["font"]
        style.font.size = Pt(size)
        if bold is not None:
            style.font.bold = bold
        rpr = style.element.get_or_add_rPr()
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rpr.append(fonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            fonts.set(qn(attr), t["font"])

    def _bottom_rule(p, color, size):
        ppr = p._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        for k, v in (("w:val", "single"), ("w:sz", str(size)), ("w:space", "1"), ("w:color", color)):
            bottom.set(qn(k), v)
        border.append(bottom)
        ppr.append(border)

    def _spacing(run, twentieths):
        if twentieths:
            rpr = run._r.get_or_add_rPr()
            sp = OxmlElement("w:spacing")
            sp.set(qn("w:val"), str(twentieths))
            rpr.append(sp)

    normal = doc.styles["Normal"]
    _font(normal, t["body"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    bullets_style = doc.styles["List Bullet"]
    _font(bullets_style, t["body"])
    bullets_style.paragraph_format.space_after = Pt(0)
    bullets_style.paragraph_format.left_indent = Inches(0.22)
    bullets_style.paragraph_format.first_line_indent = Inches(-0.16)

    def heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(t["heading_size"])
        if t["heading_color"]:
            run.font.color.rgb = _rgb(t["heading_color"])
        _spacing(run, t["heading_spacing"])
        p.paragraph_format.space_before = Pt(t["heading_before"])
        p.paragraph_format.space_after = Pt(2)
        if t["rule"]:
            _bottom_rule(p, t["rule_color"], t["rule_size"])

    name = doc.add_paragraph()
    r = name.add_run(header.get("name", ""))
    r.bold = True
    r.font.size = Pt(t["name"])
    if t["name_color"]:
        r.font.color.rgb = _rgb(t["name_color"])
    contact = doc.add_paragraph()
    rc = contact.add_run(_contact_line(header))
    rc.font.size = Pt(t["contact"])
    if t["contact_color"]:
        rc.font.color.rgb = _rgb(t["contact_color"])
    contact.paragraph_format.space_after = Pt(4 if t["name_rule"] else 2)
    if t["name_rule"]:
        _bottom_rule(contact, t["rule_color"], t["rule_size"])

    heading("Summary")
    doc.add_paragraph(content.get("summary", ""))

    heading("Experience")
    for e in content.get("experience", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.3), WD_TAB_ALIGNMENT.RIGHT)
        first, second = ((e["employer"], e["title"]) if t["role_order"] == "employer"
                         else (e["title"], e["employer"]))
        r1 = p.add_run(first)
        r1.bold = True
        r2 = p.add_run(f"  |  {second}")
        if t["dates_color"]:
            r2.font.color.rgb = _rgb(t["dates_color"])
        r3 = p.add_run(f"\t{e['dates']}")
        r3.italic = t["dates_italic"]
        if t["dates_color"]:
            r3.font.color.rgb = _rgb(t["dates_color"])
        for b in e["bullets"]:
            doc.add_paragraph(b, style="List Bullet")

    if content.get("projects"):
        heading("Selected projects")
        for pj in content["projects"]:
            p = doc.add_paragraph(style="List Bullet")
            rn = p.add_run(pj["name"])
            rn.bold = True
            p.add_run(f": {pj['line']}")

    if content.get("education"):
        heading("Education")
        for line in content["education"]:
            doc.add_paragraph(line)
    if content.get("certifications"):
        heading("Certifications")
        for line in content["certifications"]:
            doc.add_paragraph(line)
    if content.get("skills"):
        heading("Skills")
        for c in content["skills"]:
            p = doc.add_paragraph()
            rl = p.add_run(f"{c['category']}: ")
            rl.bold = True
            if t["label_color"]:
                rl.font.color.rgb = _rgb(t["label_color"])
            p.add_run(", ".join(c["items"]))

    doc.core_properties.author = header.get("name", "")
    doc.core_properties.title = (f"{header.get('name', '')} resume"
                                 + (f" for {job.title} at {job.company}" if job else ""))
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def to_pdf(docx_path: Path) -> Path | None:
    """The DOCX converted by LibreOffice; None when soffice is unavailable."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir",
                    str(docx_path.parent), str(docx_path)],
                   capture_output=True, timeout=180, check=False)
    pdf = docx_path.with_suffix(".pdf")
    return pdf if pdf.exists() else None


def page_count(pdf_path: Path) -> int:
    return len(re.findall(rb"/Type\s*/Page(?!s)", pdf_path.read_bytes()))


# ------------------------------------------------------------- pipeline

@dataclass
class DraftResult:
    docx_path: Path
    md_path: Path
    pdf_path: Path | None
    lead_role: str
    gaps: list[str] = field(default_factory=list)
    omitted_requirements: list[str] = field(default_factory=list)
    keywords: list[str] = field(default_factory=list)
    covered: list[str] = field(default_factory=list)
    missing: list[str] = field(default_factory=list)
    guard_removals: list[str] = field(default_factory=list)
    style_flags: list[str] = field(default_factory=list)
    auto_fixes: list[str] = field(default_factory=list)
    trimmed: list[str] = field(default_factory=list)
    revision_notes: list[str] = field(default_factory=list)
    title_only: bool = False
    pages: int | None = None
    word_count: int = 0


def _job_text(job: Job, title_only: bool) -> str:
    body = ("(no posting text available: draft from the title and company only)"
            if title_only else role_excerpt(job.description, 6000))
    return (f"title: {job.title}\ncompany: {job.company}\nlocation: {job.location}\n"
            f"description: {body}")


def _slug(job: Job) -> str:
    return re.sub(r"[^a-z0-9]+", "-", f"{job.company}-{job.title}".lower()).strip("-")[:80]


def _unique_stem(stem: str) -> str:
    candidate, n = stem, 1
    while (DRAFTS_DIR / f"{candidate}.docx").exists() or (DRAFTS_DIR / f"{candidate}.md").exists():
        n += 1
        candidate = f"{stem}-{n}"
    return candidate


def draft(job: Job, library: str, notes: str = "", config: dict | None = None,
          profile_text: str | None = None, style_text: str | None = None) -> DraftResult:
    """The whole thing: draft, keywords, revise, guard, lint, fit, render."""
    profile_text = PROFILE.read_text() if profile_text is None and PROFILE.exists() else (profile_text or "")
    style_text = STYLE.read_text() if style_text is None and STYLE.exists() else (style_text or "")
    header = header_from_library(library)
    title_only = len((job.description or "").strip()) < 300
    posting = _job_text(job, title_only)
    notes_block = (f"\nREQUESTER EMPHASIS NOTES (honor these):\n{notes.strip()}\n"
                   if notes and notes.strip() else "")

    raw = triage._run_claude(DRAFT_PROMPT.format(
        schema=SCHEMA, hard_rules=HARD_RULES, style=style_text, notes=notes_block,
        thesis=thesis_from_profile(profile_text), framing=framing_from_library(library),
        job=posting, library=library), DRAFT_MODEL)
    content = normalize(parse_object(raw))

    keywords: list[str] = []
    if not title_only:
        try:
            terms = triage._parse_json(triage._run_claude(
                KEYWORD_PROMPT.format(job=posting), KEYWORD_MODEL))
            keywords = clean_keywords([t for t in terms if isinstance(t, str)], posting)
        except Exception as e:  # noqa: BLE001 - keywords are best-effort
            log.warning("keyword extraction failed (%s); using fallback terms", e)
        if not keywords:
            keywords = fallback_keywords(posting, config)

    _, missing = coverage(content, keywords)
    violations = style_lint(content)
    revision_notes: list[str] = []
    if missing or violations:
        try:
            revised = normalize(parse_object(triage._run_claude(REVISE_PROMPT.format(
                missing=", ".join(missing) or "(none)",
                violations="; ".join(violations) or "(none)",
                hard_rules=HARD_RULES, style=style_text, library=library,
                draft=json.dumps(content, indent=1)), DRAFT_MODEL)))
            revision_notes = revised.pop("revision_notes", [])
            revised["gaps"] = revised["gaps"] or content["gaps"]
            revised["omitted_requirements"] = (revised["omitted_requirements"]
                                               or content["omitted_requirements"])
            content = revised
        except Exception as e:  # noqa: BLE001 - keep the first draft
            log.warning("revision pass failed (%s); keeping the first draft", e)

    removals = fabrication_guard(content, library)
    fixes = auto_fix(content)
    trimmed = fit_to_page(content)

    date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    stem = _unique_stem(f"{date}-{_slug(job)}")
    docx_path = DRAFTS_DIR / f"{stem}.docx"
    md_path = DRAFTS_DIR / f"{stem}.md"
    pdf_path, pages = None, None
    for _ in range(MAX_PDF_ROUNDS):
        render_docx(content, header, job, docx_path,
                    theme=(config or {}).get("resume_theme", DEFAULT_THEME))
        pdf_path = to_pdf(docx_path)
        pages = page_count(pdf_path) if pdf_path else None
        if pages is None or pages <= 1:
            break
        removed = trim_one(content)
        if removed is None:
            break
        trimmed.append(removed)
    md_path.write_text(render_markdown(content, header, job))

    covered, missing = coverage(content, keywords)
    return DraftResult(
        docx_path=docx_path, md_path=md_path, pdf_path=pdf_path,
        lead_role=content["lead_role"], gaps=content["gaps"],
        omitted_requirements=content["omitted_requirements"],
        keywords=keywords, covered=covered, missing=missing,
        guard_removals=removals, style_flags=style_lint(content), auto_fixes=fixes,
        trimmed=trimmed, revision_notes=revision_notes, title_only=title_only,
        pages=pages, word_count=word_count(content),
    )


# ------------------------------------------------------------------ CLI

def _main(argv: list[str]) -> int:
    """render <draft.json> <out_stem> [theme] | check <draft.json>  (no model calls)"""
    import sys
    if len(argv) < 2 or argv[0] not in ("render", "check"):
        print(_main.__doc__)
        return 2
    content = normalize(json.loads(Path(argv[1]).read_text()))
    library = LIBRARY.read_text()
    header = header_from_library(library)
    if argv[0] == "check":
        print("style:", *style_lint(content), sep="\n  ")
        print("guard:", *fabrication_guard(content, library), sep="\n  ")
        print("lines ~", estimate_lines(content), "| words", word_count(content))
        return 0
    stem = Path(argv[2])
    trimmed = fit_to_page(content)
    theme = argv[3] if len(argv) > 3 else DEFAULT_THEME
    docx_path = render_docx(content, header, None, stem.with_suffix(".docx"), theme=theme)
    stem.with_suffix(".md").write_text(render_markdown(content, header))
    pdf = to_pdf(docx_path)
    print(f"wrote {docx_path} and {stem.with_suffix('.md')}"
          + (f"; pdf {pdf} ({page_count(pdf)} page(s))" if pdf else "; no soffice, no pdf"))
    print("lines ~", estimate_lines(content), "| words", word_count(content),
          "| trimmed:", trimmed or "nothing")
    return 0


if __name__ == "__main__":
    import sys
    sys.exit(_main(sys.argv[1:]))
